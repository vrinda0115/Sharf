from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "documents"
ASSETS = OUT / "assets"
DOCX_PATH = OUT / "sharf_decision_support_blueprint.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def style_table(table, widths_dxa=None):
    table.style = "Table Grid"
    if widths_dxa:
        set_table_widths(table, widths_dxa)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    style_table(table, widths_dxa)
    return table


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def font():
    candidates = [
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, 22)
    return ImageFont.load_default()


def small_font():
    candidates = [
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, 16)
    return ImageFont.load_default()


def draw_box(draw, xy, label, fnt, fill="#F7F7F7", outline="#555555", align="center"):
    draw.rectangle(xy, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    words = label.split()
    lines = []
    line = ""
    max_chars = max(10, int((x2 - x1) / 10))
    for word in words:
        if len(line + " " + word) > max_chars:
            lines.append(line.strip())
            line = word
        else:
            line += " " + word
    if line:
        lines.append(line.strip())
    total_h = len(lines) * 22
    y = y1 + ((y2 - y1 - total_h) / 2 if align == "center" else 8)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = x1 + ((x2 - x1 - (bbox[2] - bbox[0])) / 2 if align == "center" else 10)
        draw.text((x, y), line, fill="#111111", font=fnt)
        y += 22


def arrow(draw, start, end):
    draw.line([start, end], fill="#333333", width=3)
    x1, y1 = start
    x2, y2 = end
    if x2 > x1:
        pts = [(x2, y2), (x2 - 12, y2 - 7), (x2 - 12, y2 + 7)]
    elif y2 > y1:
        pts = [(x2, y2), (x2 - 7, y2 - 12), (x2 + 7, y2 - 12)]
    else:
        pts = [(x2, y2), (x2 + 12, y2 - 7), (x2 + 12, y2 + 7)]
    draw.polygon(pts, fill="#333333")


def make_wireframes():
    ASSETS.mkdir(parents=True, exist_ok=True)
    f = font()
    sf = small_font()
    specs = []

    # Workflow diagram
    img = Image.new("RGB", (1400, 360), "white")
    d = ImageDraw.Draw(img)
    steps = [
        ("Monthly founder input", 30),
        ("Score and flag risks", 290),
        ("Compare decision options", 550),
        ("Select action plan", 810),
        ("Record outcome next month", 1070),
    ]
    for label, x in steps:
        draw_box(d, (x, 95, x + 210, 220), label, sf)
    for x in [240, 500, 760, 1020]:
        arrow(d, (x, 157), (x + 45, 157))
    d.text((30, 35), "One monthly SHARF decision cycle", fill="#111111", font=f)
    path = ASSETS / "workflow_cycle.png"
    img.save(path)
    specs.append(path)

    # Wireframe 1
    img = Image.new("RGB", (1200, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((40, 25), "Wireframe 1: Monthly Input Review", fill="#111111", font=f)
    draw_box(d, (40, 80, 1160, 140), "Purpose: confirm or edit monthly startup inputs before scoring", sf, "#EEEEEE")
    labels = ["Revenue", "MoM Growth", "Cash", "Burn", "Churn", "CAC", "Payback", "NPS", "Headcount", "Time to Fill", "Board?"]
    x, y = 60, 180
    for i, label in enumerate(labels):
        draw_box(d, (x, y, x + 250, y + 70), label + " input", sf)
        x += 280
        if x > 900:
            x = 60
            y += 95
    draw_box(d, (60, 560, 420, 640), "Run scoring", sf, "#DDDDDD")
    draw_box(d, (470, 560, 1120, 640), "Missing or high-risk fields listed here", sf)
    path = ASSETS / "wireframe_1_inputs.png"
    img.save(path)
    specs.append(path)

    # Wireframe 2
    img = Image.new("RGB", (1200, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((40, 25), "Wireframe 2: Health Score and Risk Explanation", fill="#111111", font=f)
    draw_box(d, (40, 80, 330, 260), "Composite SHARF Score 0-100", sf, "#EEEEEE")
    draw_box(d, (370, 80, 760, 260), "Risk Tier Low Watch Critical", sf)
    draw_box(d, (800, 80, 1160, 260), "Plain-language reason", sf)
    dims = ["Runway", "Revenue", "Retention", "Growth Efficiency", "Customer Sentiment", "Team Capacity", "Governance"]
    y = 310
    for dim in dims:
        draw_box(d, (70, y, 320, y + 45), dim, sf)
        d.rectangle((350, y + 10, 850, y + 35), outline="#555555", fill="#F7F7F7", width=1)
        d.rectangle((350, y + 10, 350 + (dims.index(dim) + 2) * 45, y + 35), fill="#999999")
        draw_box(d, (900, y, 1130, y + 45), "Why it matters", sf)
        y += 55
    path = ASSETS / "wireframe_2_score.png"
    img.save(path)
    specs.append(path)

    # Wireframe 3
    img = Image.new("RGB", (1200, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((40, 25), "Wireframe 3: Decision Options", fill="#111111", font=f)
    for i, title in enumerate(["Conserve Cash", "Maintain Plan", "Prepare Fundraise"]):
        x = 50 + i * 380
        draw_box(d, (x, 90, x + 340, 170), title, f, "#EEEEEE")
        draw_box(d, (x, 195, x + 340, 290), "Expected effect on runway", sf)
        draw_box(d, (x, 310, x + 340, 405), "Tradeoffs and risk flags", sf)
        draw_box(d, (x, 425, x + 340, 520), "Recommended next action", sf)
    draw_box(d, (70, 590, 1130, 670), "User selects option, adds note, and exports founder action plan", sf, "#DDDDDD")
    path = ASSETS / "wireframe_3_options.png"
    img.save(path)
    specs.append(path)

    # Wireframe 4
    img = Image.new("RGB", (1200, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((40, 25), "Wireframe 4: Action Plan and Outcome Log", fill="#111111", font=f)
    draw_box(d, (50, 90, 1150, 150), "Purpose: convert recommendation into a 30-day operating decision", sf, "#EEEEEE")
    rows = ["Decision", "Owner", "Metric Target", "Due Date", "Status"]
    for i, h in enumerate(rows):
        draw_box(d, (70 + i * 215, 200, 260 + i * 215, 250), h, sf, "#DDDDDD")
        draw_box(d, (70 + i * 215, 255, 260 + i * 215, 350), "entry", sf)
        draw_box(d, (70 + i * 215, 355, 260 + i * 215, 450), "entry", sf)
    draw_box(d, (70, 520, 560, 650), "Next month outcome notes", sf)
    draw_box(d, (620, 520, 1130, 650), "Did score improve? Yes / No / Review", sf)
    path = ASSETS / "wireframe_4_action_log.png"
    img.save(path)
    specs.append(path)
    return specs


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.1
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(46, 116, 181)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor(46, 116, 181)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(31, 77, 120)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)


def build_doc():
    OUT.mkdir(parents=True, exist_ok=True)
    wireframes = make_wireframes()
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SHARF Decision-Support Tool Blueprint and Scope Check")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Startup health and risk forecasting | Conceptual design assignment")
    add_para(doc, "Prepared for a decision-support prototype focused on early-stage founder decisions about burn rate, fundraising readiness, and operational remediation.")

    doc.add_heading("1. Decision recap", level=1)
    add_para(doc, "Decision-maker: Non-technical early-stage startup founder, typically seed to Series A, who tracks basic performance metrics but does not have a dedicated data or finance team.", bold_prefix="Decision-maker:")
    add_para(doc, "Problem to be addressed: Founders often miss slow-moving financial and operational warning signs because raw metrics are reviewed separately and informally. SHARF converts founder-reported inputs into a composite health score, dimension-level risk signals, and plain-language decision guidance.", bold_prefix="Problem to be addressed:")
    add_para(doc, "Key decisions: The tool supports three recurring decisions: whether to reduce or sustain burn, whether the startup is ready to approach investors, and which operational dimension should receive the next 30-day remediation focus.", bold_prefix="Key decisions:")
    add_para(doc, "Decision frequency: Monthly for burn and operational remediation decisions; quarterly for fundraising readiness review.", bold_prefix="Decision frequency:")
    add_para(doc, "Decision time horizon: 30-90 days for operational response, 6-18 months for fundraising readiness, and 3-5 years for survival-risk framing.", bold_prefix="Decision time horizon:")
    add_para(doc, "Most important metrics: runway in months, month-over-month revenue growth, churn, CAC payback, NPS or customer satisfaction, headcount, average time to fill roles, governance structure, composite SHARF score, dimension scores, and predicted survival probability.", bold_prefix="Most important metrics:")
    add_para(doc, "Framing change: The prior framing remains SHARF as a founder-facing startup health and risk tool. This blueprint narrows the course MVP to monthly decision support rather than a live investor-grade forecasting platform.", bold_prefix="Framing change:")

    doc.add_heading("2. Decision logic specification", level=1)
    add_para(doc, "SHARF uses a transparent AHP-SAW logic: normalize each metric into a 0-100 dimension score, apply AHP-derived weights, sum weighted scores into a composite health score, classify risk tiers, and expose the reason for each recommendation.")
    rows = [
        ("Burn rate decision", "Cash balance, monthly burn, runway, MoM revenue growth, macro stress", "Runway = cash / monthly burn; trend check uses 3-month revenue growth and burn change", "Critical if runway < 6 months; watch if 6-12 months; stable if > 12 months. If growth is negative and burn rising, recommend immediate burn review.", "Reduce burn, hold burn, or approve limited spend with target runway."),
        ("Fundraising readiness", "Composite SHARF score, runway, growth, churn, CAC payback, NPS, governance", "Weighted score plus dimension gates for investor credibility", "Ready if composite >= 75, runway >= 12 months, churn not rising, CAC payback acceptable, and board/advisory oversight present. Prepare if 60-74 or one gate weak. Delay if < 60 or runway < 6.", "Approach investors now, prepare in 60-90 days, or delay and fix specific blockers."),
        ("Operational remediation priority", "Dimension scores for retention, growth efficiency, sentiment, team capacity, governance", "Lowest weighted dimension and largest negative trend are ranked first", "Prioritize the dimension with score < 60 and highest AHP weight; break ties by fastest deterioration over 3 months.", "30-day action focus with metric target and owner."),
        ("Survival-risk interpretation", "Composite score, BLS-calibrated survival probability, macro stress, stage", "Logistic model output calibrated to five-year survival base rates", "Do not present as certainty; flag as context if survival probability falls below scenario benchmark.", "Plain-language risk explanation and scenario review rather than deterministic prediction."),
    ]
    add_table(doc, ["Decision", "Inputs reviewed", "Metric/model output", "Rules and thresholds", "Recommendation presented"], rows, [1500, 2100, 2000, 2200, 1560])
    add_para(doc, "Structured pseudo-code:")
    add_numbered(doc, [
        "Collect current monthly startup snapshot and validate missing or implausible entries.",
        "Compute runway, MoM revenue growth, churn, CAC payback, NPS score, hiring capacity score, and governance score.",
        "Normalize all metrics so higher values consistently mean healthier performance.",
        "Apply AHP weights and SAW aggregation to calculate composite SHARF score and dimension scores.",
        "Apply decision gates for burn, fundraising, and remediation priority.",
        "Return recommendation, supporting reasons, and the next action log fields."
    ])

    doc.add_heading("3. Primary user and use cases", level=1)
    add_para(doc, "Primary user: A non-technical early-stage founder who must make monthly operating decisions without a formal finance function.", bold_prefix="Primary user:")
    add_para(doc, "Secondary user: A mentor, advisor, accelerator coach, or board observer who reviews the founder's monthly decision rationale and helps pressure-test assumptions.", bold_prefix="Secondary user:")
    rows = [
        ("Monthly runway check", "Founder closes monthly books and updates cash and burn.", "Whether to reduce burn or continue planned spend.", "Runway, burn trend, revenue growth, macro stress, recommendation reason.", "Approve a spend posture and record target runway for next month."),
        ("Investor outreach review", "Founder is considering seed or Series A outreach this quarter.", "Whether to approach investors now or delay.", "Composite score, fundraising gates, weak dimensions, governance signal.", "Start outreach, prepare materials, or fix blockers first."),
        ("Retention warning", "Churn rises for two months while revenue growth slows.", "Whether retention should override acquisition spending.", "Churn trend, NPS, revenue growth, CAC payback interaction.", "Prioritize retention action and pause inefficient acquisition spend."),
        ("Hiring capacity review", "Open roles are staying unfilled and delivery milestones are slipping.", "Whether hiring process risk is constraining growth.", "Headcount, time to fill, macro labor context, team-capacity score.", "Adjust hiring plan, contractor strategy, or milestone commitments."),
        ("Advisor check-in", "Monthly advisor meeting requires evidence behind founder decisions.", "Whether the plan is defensible and transparent.", "Inputs, thresholds, score history, selected action and outcome log.", "Review rationale and agree on next-month accountability metrics."),
    ]
    add_table(doc, ["Use case", "Situation or trigger", "Decision needed", "Information needed", "Expected action"], rows, [1500, 2100, 1900, 2100, 1760])

    doc.add_heading("4. Workflow across one decision cycle", level=1)
    add_para(doc, "The decision cycle is monthly because the core inputs, including revenue, burn, churn, cash balance, CAC payback, NPS, and hiring status, are most actionable at monthly granularity. The founder opens SHARF after monthly financial and operating numbers are available, enters or reviews the current snapshot, adjusts optional assumptions, reviews the score and explanations, compares decision options, and records the action selected.")
    doc.add_picture(str(wireframes[0]), width=Inches(6.5))
    add_numbered(doc, [
        "Access: Founder opens SHARF during the monthly operating review, ideally before advisor or board check-in.",
        "Inputs: Founder reviews cash, burn, revenue, churn, CAC, payback, NPS, headcount, time to fill, and governance status.",
        "Adjustable settings: Stage, risk tolerance, target runway, and optional scenario assumptions such as burn reduction percentage.",
        "Outputs: Composite score, dimension scores, risk tier, recommendations, explanations, and recommended 30-day actions.",
        "Comparison: Founder compares conserve cash, maintain plan, and prepare fundraise options using runway and risk impacts.",
        "Action: Founder selects a decision, assigns an owner or follow-up, and records next month's target metric.",
        "Review: At the next cycle, the outcome log is compared with the new score to determine whether the action improved the risk profile."
    ])

    doc.add_heading("5. Low-fidelity wireframes", level=1)
    add_para(doc, "The wireframes are intentionally grayscale and low fidelity. They show structure, decision flow, and explanation placement rather than visual polish.")
    captions = [
        "Screen purpose: confirm the monthly data snapshot before scoring. Decision step supported: determine whether the score is based on complete and current inputs.",
        "Screen purpose: review SHARF score, dimension scores, and explanation. Decision step supported: identify the risk tier and the reason behind it.",
        "Screen purpose: compare recommended decision options. Decision step supported: choose between conserving cash, maintaining the plan, or preparing to fundraise.",
        "Screen purpose: record the selected action and next-month outcome. Decision step supported: convert recommendation into accountable follow-up."
    ]
    for path, caption in zip(wireframes[1:], captions):
        doc.add_picture(str(path), width=Inches(6.5))
        p = add_para(doc, caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("6. Feature-to-decision traceability table", level=1)
    rows = [
        ("Monthly input form", "All decisions", "Founder-entered monthly snapshot fields", "Confirms the evidence base before scoring", "Prevents recommendations from appearing detached from data."),
        ("Runway calculator", "Burn rate decision", "Cash balance and monthly burn", "Shows how many months remain before cash runs out", "Runway is the main early warning signal for distress."),
        ("Revenue trajectory panel", "Burn and fundraising decisions", "Monthly revenue and MoM growth", "Shows whether growth supports current spend and investor timing", "Revenue trajectory is a strong predictor of financial outcomes."),
        ("Retention and sentiment panel", "Remediation priority", "Churn and NPS", "Shows whether demand-side risk is emerging before revenue decline", "Retention problems often precede weaker revenue."),
        ("Growth efficiency panel", "Spend and fundraising decisions", "CAC and CAC payback", "Shows whether acquisition spending is sustainable", "High CAC payback can make growth unhealthy even when revenue rises."),
        ("Team capacity panel", "Operational remediation", "Headcount and average time to fill roles", "Shows whether hiring constraints threaten execution", "Team composition and hiring velocity affect fundraising credibility."),
        ("Governance flag", "Fundraising readiness", "Board or advisory group present", "Shows whether accountability structure is credible", "Investors value structured oversight signals."),
        ("Recommendation explanation", "All decisions", "Thresholds, score drivers, and trend logic", "Shows why the tool suggests an action", "Decision support requires transparent reasoning, not just a score."),
        ("Action log", "Decision follow-through", "Selected action, owner, target, outcome", "Turns recommendation into a reviewable decision record", "Makes monthly learning and accountability possible."),
    ]
    add_table(doc, ["Tool component or screen element", "Decision supported", "Metric, model, or input used", "What the user learns or does", "Why necessary"], rows, [1900, 1600, 2000, 2100, 1760])

    doc.add_heading("7. MVP scope and out-of-scope statement", level=1)
    add_para(doc, "MVP scope: The course MVP will support one founder-facing monthly decision cycle for a seed-to-Series A startup. It will accept or simulate the required monthly inputs, compute core metrics, normalize dimension scores, calculate a composite SHARF score using fixed AHP-SAW weights, classify risk tiers, display explanatory decision logic, and produce a 30-day action recommendation.")
    add_para(doc, "The MVP will include these data inputs: monthly revenue and MoM growth; cash balance and monthly burn rate; monthly churn; CAC and CAC payback; NPS or equivalent customer satisfaction score; headcount and average time to fill roles; and a binary governance indicator for board or advisory oversight.")
    add_para(doc, "Out of scope: Live accounting integrations, bank connections, automated CRM ingestion, investor matching, legally binding fundraising advice, real-time forecasting, industry-specific benchmark subscriptions, multi-user permissions, and production security workflows.")
    add_para(doc, "Why realistic: The required inputs can be entered manually or generated synthetically. The scoring logic is transparent, the prototype dataset can remain small, and the main contribution is decision structure rather than enterprise-grade automation.")
    add_para(doc, "Stretch goals: Scenario sliders for burn reduction, exportable advisor summary, additional macro-context indicators, industry-specific AHP weights, and historical score trend charts.")

    doc.add_heading("8. Evaluation plan", level=1)
    add_para(doc, "Functional evaluation will test whether the tool implements the intended logic correctly. Test cases will include normal, watch, and critical scenarios for runway, revenue growth, churn, CAC payback, NPS, hiring velocity, and governance. Each test case will specify expected metric values, normalized dimension scores, composite score range, risk tier, and recommendation.")
    add_bullets(doc, [
        "Metric tests: verify formulas for runway, MoM growth, churn, CAC payback, and rolling trends.",
        "Threshold tests: verify that runway below 6 months, worsening churn, and weak CAC payback trigger the correct flags.",
        "Scoring tests: verify that AHP-SAW weights sum correctly and that higher-risk inputs lower the composite score.",
        "Data-quality tests: verify null checks, impossible values, and range checks for churn, NPS, revenue, and burn.",
        "Traceability tests: verify that every recommendation displays the score drivers and rules behind it."
    ])
    add_para(doc, "Decision-support evaluation will assess whether SHARF helps the intended founder make a clearer, faster, and more consistent decision. The evaluation will use scenario walkthroughs with a founder or advisor persona, think-aloud review of the recommendation screens, and comparison against an unstructured spreadsheet review.")
    add_bullets(doc, [
        "Comprehensibility: Can the founder explain why SHARF made the recommendation?",
        "Workflow fit: Does the monthly flow match when the founder actually has the data?",
        "Actionability: Does the tool produce a concrete next step rather than a passive dashboard observation?",
        "Consistency: Do similar startup scenarios lead to similar decisions across repeated reviews?",
        "Usefulness: Does the founder or advisor believe the tool would improve the operating review conversation?"
    ])

    doc.add_heading("References", level=1)
    refs = [
        "Abdullah, N., et al. (2024). A multi-criteria scoring system for small business eligibility assessment. Journal of Small Business Management, 62(3), 45-67.",
        "Federal Reserve Bank of St. Louis. (2025). Federal Funds Effective Rate [FEDFUNDS], CPIAUCSL, Business Applications [BUSAPPWNSACO], and Unemployment Rate [UNRATE]. FRED Economic Data.",
        "Mello, A., et al. (2026). Fundraising success pathways in early-stage startups: A multi-type analysis. Entrepreneurship Theory and Practice, 50(1), 112-138.",
        "Mseddi, S. (2026). Predicting startup financial outcomes using operational indicators: Evidence from 4,400 ventures. Strategic Management Journal, 47(2), 201-228.",
        "Oliva, F. L., et al. (2022). Structured multi-dimensional risk frameworks for early-stage ventures. Journal of Business Venturing, 37(4), 106-123.",
        "Rady, M., et al. (2025). Early warning indicators and investor communication in startup distress. Journal of Finance and Entrepreneurship, 18(1), 55-78.",
        "Safari, H., et al. (2024). Operational risk categorization in small and early-stage businesses. International Journal of Operations & Production Management, 44(3), 310-334.",
        "Statista Market Insights. (2026). Venture capital - United States: Capital raised by stage, 2017-2026.",
        "U.S. Bureau of Labor Statistics. (2025). Business Employment Dynamics Tables 5 and 7."
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(ref)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
