from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FINAL_DIR = ROOT / "data" / "final"
OUTPUT_PATH = ROOT / "predictive_model_decision_report.docx"


def set_run_font(run, size: float = 9.5, bold: bool = False, color: str = "102033") -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_label_paragraph(document: Document, label: str, body: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.02
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run)


def main() -> None:
    model = json.loads((FINAL_DIR / "runway_prediction_model.json").read_text(encoding="utf-8"))
    metrics = model["metrics"]

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("SHARF Predictive Model Decision Memo")
    set_run_font(title_run, size=15.5, bold=True, color="092C56")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(7)
    subtitle_run = subtitle.add_run("Three-month runway forecasting and what-if analysis for founder cash decisions")
    set_run_font(subtitle_run, size=8.8, color="668CA9")

    add_label_paragraph(
        document,
        "Decision-maker need. ",
        "The SHARF decision-maker is a founder deciding whether to conserve cash, maintain the current operating plan, or prepare for fundraising. The predictive need is not a generic score; it is an estimate of runway three months ahead under current or adjusted operating conditions, plus a clear statement of uncertainty.",
    )
    add_label_paragraph(
        document,
        "Modeling approach. ",
        "I used ridge regression to predict future runway months from ten monthly predictors: revenue, cash, burn, revenue growth, churn, CAC payback, NPS, headcount, hiring time, and governance score. Ridge regression is appropriate for this prototype because the dataset is small, the predictors are correlated, and the decision-maker benefits from a stable, transparent model rather than a harder-to-explain black-box method.",
    )
    add_label_paragraph(
        document,
        "Validation design. ",
        f"The model was trained only on earlier startup-months and tested on later target months beginning {model['test_target_start']}. The temporal holdout contained {model['test_rows']} rows, which protects the evaluation from using future outcomes during fitting.",
    )

    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["MAE", "RMSE", "R^2", "90% PI coverage"]
    values = [
        f"{metrics['mae']:.2f} months",
        f"{metrics['rmse']:.2f} months",
        f"{metrics['r2']:.3f}",
        f"{metrics['interval_coverage_90']:.1%}",
    ]
    for cell, value in zip(table.rows[0].cells, headers):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        set_run_font(run, size=8.6, bold=True, color="092C56")
    for cell, value in zip(table.rows[1].cells, values):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        set_run_font(run, size=8.6)

    add_label_paragraph(
        document,
        "Dashboard outputs. ",
        "The Predictive Model Performance screen shows the holdout error metrics, interval coverage, and an actual-versus-predicted chart so the founder can judge whether the model is directionally useful before relying on it. The What-If Runway Analysis screen lets the founder modify feature values and immediately receives a point prediction with a 90% predictive interval.",
    )
    add_label_paragraph(
        document,
        "How uncertainty informs action. ",
        "The interval is used as the decision guardrail. If the lower bound falls below six months, SHARF flags the scenario for cash conservation or earlier fundraising preparation. If the point estimate is acceptable but the interval crosses the risk threshold, the founder can treat the result as ambiguous and avoid overcommitting to a single plan.",
    )
    add_label_paragraph(
        document,
        "Limitations. ",
        "The current data are synthetic and validated for prototype testing, not production forecasting. The interval reflects model residuals and unusual feature combinations in this dataset; real deployment would require retraining on historical startup operating data, monitoring calibration, and refreshing the model as market conditions change.",
    )

    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
