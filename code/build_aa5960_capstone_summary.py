from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FINAL_DIR = ROOT / "data" / "final"
OUTPUT_DIR = ROOT / "output" / "aa5960_report_assets"
DOCX_PATH = ROOT / "AA5960_SHARF_Capstone_Project_Summary_Detailed.docx"


BLUE = "2E5E97"
INK = "000000"
MUTED = "4D4D4D"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/times.ttf"),
        Path("C:/Windows/Fonts/timesbd.ttf") if bold else Path("C:/Windows/Fonts/times.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def set_run_font(run, size: float = 12, bold: bool = False, italic: bool = False, color: str = INK) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before: float = 0, after: float = 0, line: float = 1.08) -> None:
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_numbered_heading(document: Document, number: int, title: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=4)
    paragraph.paragraph_format.left_indent = Inches(0.0)
    paragraph.paragraph_format.first_line_indent = Inches(0.0)
    run = paragraph.add_run(f"{number}.    {title}")
    set_run_font(run, bold=True)


def add_inline_label_paragraph(document: Document, label: str, body: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph(paragraph)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run)


def add_reference(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=1, line=0.95)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, size: float = 10.3, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, alignment=align, after=0, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_report_table(document: Document, headers: list[str], rows: list[list[str]], caption: str, widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        shade_cell(cell, "E8EEF5")
        set_cell(cell, header, size=10.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            set_cell(cells[idx], value, size=10.1)
    cap = document.add_paragraph()
    set_paragraph(cap, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.0)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=11, italic=True)


def draw_trend_chart(metrics: pd.DataFrame) -> Path:
    path = OUTPUT_DIR / "figure_1_monthly_trends.png"
    frame = metrics.copy()
    frame["month"] = pd.to_datetime(frame["month"])
    monthly = frame.groupby("month")[["sharf_score", "runway_months"]].mean().reset_index()
    img = Image.new("RGB", (1000, 560), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(30, True)
    text_font = font(20)
    small_font = font(17)
    draw.text((40, 28), "Average SHARF Score and Runway by Month", fill="#092C56", font=title_font)
    draw.line((575, 70, 625, 70), fill="#225688", width=4)
    draw.text((635, 58), "SHARF score", fill="#225688", font=small_font)
    draw.line((760, 70, 810, 70), fill="#1F7A5A", width=4)
    draw.text((820, 58), "Runway months", fill="#1F7A5A", font=small_font)
    left, top, right, bottom = 110, 105, 890, 450
    draw.rectangle((left, top, right, bottom), outline="#222222", width=2)

    months = list(monthly["month"])
    score_values = monthly["sharf_score"].tolist()
    runway_values = monthly["runway_months"].tolist()
    min_score, max_score = 0, 100
    min_runway, max_runway = 0, max(18, int(max(runway_values) + 4))

    for tick in range(0, 101, 25):
        y = bottom - (bottom - top) * tick / 100
        draw.line((left - 6, y, left, y), fill="#222222", width=1)
        draw.text((left - 48, y - 10), str(tick), fill="black", font=small_font)
    for tick in range(0, max_runway + 1, max(3, max_runway // 4)):
        y = bottom - (bottom - top) * tick / max_runway
        draw.line((right, y, right + 6, y), fill="#222222", width=1)
        draw.text((right + 12, y - 10), str(tick), fill="black", font=small_font)

    def point(index: int, value: float, value_min: float, value_max: float) -> tuple[float, float]:
        x = left + (right - left) * index / max(len(months) - 1, 1)
        y = bottom - (bottom - top) * (value - value_min) / (value_max - value_min)
        return x, y

    score_points = [point(i, v, min_score, max_score) for i, v in enumerate(score_values)]
    runway_points = [point(i, v, min_runway, max_runway) for i, v in enumerate(runway_values)]
    draw.line(score_points, fill="#225688", width=4)
    draw.line(runway_points, fill="#1F7A5A", width=4)
    for x, y in score_points:
        draw.ellipse((x - 4, y - 4, x + 4, y + 4), fill="#225688")
    for x, y in runway_points:
        draw.rectangle((x - 4, y - 4, x + 4, y + 4), fill="#1F7A5A")
    for i, month in enumerate(months):
        if i in {0, len(months) - 1, len(months) // 2}:
            x, _ = point(i, 0, 0, 1)
            draw.text((x - 26, bottom + 16), month.strftime("%b"), fill="black", font=small_font)
    draw.text((126, 470), "Left axis: score", fill="#225688", font=text_font)
    draw.text((620, 470), "Right axis: months", fill="#1F7A5A", font=text_font)
    draw.text((40, 520), "Source: SHARF validated monthly metrics, averaged across 18 startups.", fill="#4D4D4D", font=small_font)
    img.save(path)
    return path


def draw_risk_chart(metrics: pd.DataFrame) -> Path:
    counts = metrics["risk_tier"].value_counts().reindex(["Critical", "Watch", "Stable"], fill_value=0)
    path = OUTPUT_DIR / "figure_1_risk_tiers.png"
    img = Image.new("RGB", (1000, 520), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(30, True)
    text_font = font(24)
    small_font = font(20)
    colors = {"Critical": "#B42318", "Watch": "#B7791F", "Stable": "#1F7A5A"}
    draw.text((40, 30), "Risk Tier Distribution Across Startup-Months", fill="#092C56", font=title_font)
    max_count = max(counts.max(), 1)
    x0, y0 = 230, 120
    bar_h, gap = 58, 36
    for idx, (tier, count) in enumerate(counts.items()):
        y = y0 + idx * (bar_h + gap)
        width = int(650 * count / max_count)
        draw.text((40, y + 13), tier, fill="black", font=text_font)
        draw.rectangle((x0, y, x0 + 650, y + bar_h), outline="#D9D9D9", width=2)
        draw.rectangle((x0, y, x0 + width, y + bar_h), fill=colors[tier])
        draw.text((x0 + width + 16, y + 13), str(int(count)), fill="black", font=text_font)
    draw.text((40, 450), "Source: SHARF validated monthly metrics, n=216 startup-months.", fill="#4D4D4D", font=small_font)
    img.save(path)
    return path


def draw_prediction_chart(holdout: pd.DataFrame) -> Path:
    path = OUTPUT_DIR / "figure_2_prediction_holdout.png"
    img = Image.new("RGB", (1000, 620), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(30, True)
    text_font = font(21)
    small_font = font(18)
    draw.text((40, 30), "Actual vs. Predicted Three-Month Runway", fill="#092C56", font=title_font)
    left, top, right, bottom = 120, 110, 900, 510
    draw.rectangle((left, top, right, bottom), outline="#222222", width=2)
    max_val = max(float(holdout["target_runway_months_3mo"].max()), float(holdout["predicted_runway_months_3mo"].max()), 1.0)
    max_val = int(max_val + 4)
    for tick in range(0, max_val + 1, max(1, max_val // 5)):
        x = left + (right - left) * tick / max_val
        y = bottom - (bottom - top) * tick / max_val
        draw.line((x, bottom, x, bottom + 8), fill="#222222", width=1)
        draw.line((left - 8, y, left, y), fill="#222222", width=1)
        draw.text((x - 12, bottom + 14), str(tick), fill="black", font=small_font)
        draw.text((left - 48, y - 10), str(tick), fill="black", font=small_font)
    draw.line((left, bottom, right, top), fill="#668CA9", width=3)
    for _, row in holdout.iterrows():
        actual = float(row["target_runway_months_3mo"])
        predicted = float(row["predicted_runway_months_3mo"])
        x = left + (right - left) * actual / max_val
        y = bottom - (bottom - top) * predicted / max_val
        draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill="#225688", outline="#092C56")
    draw.text((350, 555), "Actual runway three months later", fill="black", font=text_font)
    draw.text((18, 280), "Predicted", fill="black", font=text_font)
    draw.text((40, 585), "Dashed diagonal equivalent is shown as a solid reference line; closer points indicate stronger predictive accuracy.", fill="#4D4D4D", font=small_font)
    img.save(path)
    return path


def add_figure(document: Document, image_path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.9))
    cap = document.add_paragraph()
    set_paragraph(cap, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.0)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=11, italic=True)


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    metrics = pd.read_csv(FINAL_DIR / "sharf_monthly_metrics.csv")
    profiles = pd.read_csv(FINAL_DIR / "startup_profiles.csv")
    forecast = pd.read_csv(FINAL_DIR / "monte_carlo_risk_forecast.csv")
    validation = pd.read_csv(FINAL_DIR / "validation_summary.csv")
    holdout = pd.read_csv(FINAL_DIR / "runway_model_holdout_predictions.csv")
    model = json.loads((FINAL_DIR / "runway_prediction_model.json").read_text(encoding="utf-8"))
    risk_counts = metrics["risk_tier"].value_counts().to_dict()
    risk_total = len(metrics)
    interval_width = (holdout["prediction_interval_upper_90"] - holdout["prediction_interval_lower_90"]).mean()

    trend_chart = draw_trend_chart(metrics)
    risk_chart = draw_risk_chart(metrics)
    prediction_chart = draw_prediction_chart(holdout)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.different_first_page_header_footer = True

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    set_paragraph(title, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=64, after=18, line=1.0)
    title_run = title.add_run("SHARF Founder Decision Support: Capstone Project Summary")
    set_run_font(title_run, size=20, bold=True, color=BLUE)

    for line in [
        "Vrinda Linesh Thakur",
        "Saint Louis University",
        "AA 5960",
        "Capstone Project Written Summary",
        "August 6, 2026",
    ]:
        paragraph = doc.add_paragraph()
        set_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=4, line=1.0)
        run = paragraph.add_run(line)
        set_run_font(run, size=16)

    doc.add_section(WD_SECTION.NEW_PAGE)
    section2 = doc.sections[-1]
    section2.top_margin = Inches(1)
    section2.bottom_margin = Inches(1)
    section2.left_margin = Inches(1)
    section2.right_margin = Inches(1)

    heading = doc.add_paragraph()
    set_paragraph(heading, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=1.0)
    run = heading.add_run("AA 5960 Analysis: SHARF Founder Risk Cockpit")
    set_run_font(run, bold=True)

    add_numbered_heading(doc, 1, "Purpose and Scope")
    add_body_paragraph(
        doc,
        "The purpose of this project was to design and implement SHARF, a software-based analytics prototype that helps early-stage startup founders make monthly decisions about cash conservation, fundraising readiness, and operational risk. The project responds to a practical decision problem: founders often have scattered signals about revenue, burn, churn, hiring, customer sentiment, and governance, but they need a concise view of whether the company is healthy enough to continue its plan or whether immediate action is needed.",
    )
    add_body_paragraph(
        doc,
        "The scope included synthetic startup profile data, monthly operating metrics, a transparent health score, Monte Carlo risk forecasting, a supervised three-month runway prediction model, and a Streamlit application with what-if analysis. The system is not intended to replace investor diligence or financial planning. Its purpose is to demonstrate how validated data, predictive modeling, and uncertainty communication can be turned into a founder-facing decision-support workflow.",
    )

    add_numbered_heading(doc, 2, "Literature Review")
    add_body_paragraph(
        doc,
        "The literature supporting SHARF shows that startup health cannot be evaluated from one metric alone. Rady et al. (2025) analyze AI startup financing rounds and show that investor valuation is shaped not only by operating signals but also by how founders communicate credibility and comprehensibility. This finding supports SHARF's plain-language recommendation layer: the application should not merely calculate a score, but help founders translate that score into evidence-backed language for board, mentor, and fundraising conversations.",
    )
    add_body_paragraph(
        doc,
        "A second stream of literature emphasizes context and ecosystem variation. Menshikov et al. (2025) use correlation and regression models to identify ecosystem factors, such as funding access, policy support, and market size, that influence startup financing outcomes. Mseddi (2026) similarly uses Crunchbase deal data to show that funding rounds, founders, employees, investors, and funding type are significant predictors of startup financing success across MENA countries. Together, these studies justify SHARF's benchmark comparison design: founders should be compared against companies with similar stages, verticals, and financing paths rather than against a universal threshold.",
    )
    add_body_paragraph(
        doc,
        "The third stream focuses on multi-factor success pathways and risk categories. Mello et al. (2026) use fuzzy-set Qualitative Comparative Analysis to show that sustainable startups can reach crowdfunding success through different configurations of credibility, narrative, governance, and investor endorsement. Safari et al. (2024), in a systematic review of supply-chain risk and resilience, classify risks into demand, supply, organizational, operational, environmental, and network/control dimensions. Oliva et al. (2022) also organize born-global startup risk across social, environmental, economic, and institutional dimensions. These studies support SHARF's design choice to separate cash, churn, hiring, operational, governance, and customer-sentiment signals while still combining them into a founder-facing risk view.",
    )
    add_body_paragraph(
        doc,
        "The fourth stream supports the decision-support method. Abdullah et al. (2024) demonstrate that a lightweight decision support system using Analytic Hierarchy Process and Simple Additive Weighting can combine multiple financial and operational criteria into an interpretable eligibility score for small businesses. This is directly relevant to SHARF's health score, which combines revenue, burn rate, churn, headcount, CAC payback, NPS, governance, and hiring indicators into a composite view. The predictive analytics literature extends this foundation: Shmueli (2010) distinguishes prediction from explanation, Breiman (2001) emphasizes out-of-sample predictive performance, Hoerl and Kennard (1970) justify ridge regression for correlated predictors, and Goodwin et al. (2010) show why prediction intervals should be communicated alongside point forecasts. The research questions guiding this project are: Which monthly startup signals can be combined into an interpretable founder risk view, and how can predictive performance plus prediction uncertainty improve founder decisions about burn reduction and fundraising timing?",
    )

    add_numbered_heading(doc, 3, "Method")
    add_inline_label_paragraph(
        doc,
        "Data and validation. ",
        f"The project used {len(profiles)} synthetic startup profiles and {len(metrics)} startup-month records. No private founder, customer, employee, investor, or transaction data were used. The final data package includes startup profiles, monthly operating metrics, Monte Carlo forecast outputs, a sourced company employment/payroll benchmark, a data dictionary, a research evidence map, model artifacts, and machine-readable validation results. The validation script checked schemas, identifiers, ranges, missingness, formulas, referential integrity, forecast outputs, predictive-model outputs, and benchmark quality; {int(validation['passed'].astype(str).str.lower().eq('true').sum())} of {len(validation)} checks passed.",
    )
    add_report_table(
        doc,
        ["Data asset", "Rows", "Purpose in analysis"],
        [
            ["Startup profiles", f"{len(profiles)}", "Defines stage, industry, region, governance, and founder-profile context."],
            ["Monthly metrics", f"{len(metrics)}", "Main analytical table for runway, revenue growth, churn, CAC payback, NPS, hiring, and SHARF score."],
            ["Monte Carlo forecast", f"{len(forecast)}", "Summarizes six-month probabilistic risk for each startup from 1,000 simulated paths."],
            ["Model holdout predictions", f"{len(holdout)}", "Compares actual and predicted three-month runway values for unseen future months."],
        ],
        "Table 1. Final SHARF data assets used for scoring, forecasting, and model evaluation.",
        [1.55, 0.7, 4.25],
    )
    add_inline_label_paragraph(
        doc,
        "Analytics workflow. ",
        "The SHARF score was built from normalized operating dimensions: runway, revenue growth, retention, growth efficiency, customer sentiment, team capacity, and governance. A Monte Carlo layer simulated 1,000 six-month paths per startup to estimate risk probabilities. The predictive-model layer used ridge regression to forecast runway three months ahead from ten monthly predictors, then evaluated performance on a temporal holdout rather than on the same months used for training.",
    )
    add_inline_label_paragraph(
        doc,
        "Prototype implementation. ",
        "The Streamlit application contains six screens: founder input, health score and risk, decision options, action plan and outcome log, predictive model performance, and what-if runway analysis. The what-if screen allows the decision-maker to change feature values such as burn, cash, growth, churn, CAC payback, hiring time, and governance, then receive a point prediction and a 90% predictive interval.",
    )
    doc.add_page_break()
    add_report_table(
        doc,
        ["Page", "Main output", "Decision-maker use"],
        [
            ["1. Founder Input", "Editable operating metrics for revenue, cash, burn, churn, CAC payback, NPS, team, and governance.", "Lets the founder test the current month and correct assumptions before scoring."],
            ["2. Health Score and Risk", "Composite SHARF score, risk tier, dimension scores, and plain-language explanation.", "Shows which business dimension is weakest and whether the company is Stable, Watch, or Critical."],
            ["3. Decision Options", "Compares Conserve Cash, Maintain Plan, and Prepare Fundraise choices.", "Turns analytics into operating choices with expected runway effects and tradeoffs."],
            ["4. Action Plan and Outcome Log", "Decision owner, due date, status, outcome notes, and score-memory log.", "Creates a feedback loop so the founder can compare later outcomes with the earlier decision."],
            ["5. Predictive Model Performance", "Holdout MAE, RMSE, R^2, interval coverage, and actual-vs-predicted chart.", "Helps the founder judge whether the predictive model is reliable enough to support planning."],
            ["6. What-If Runway Analysis", "Scenario inputs plus predicted three-month runway and 90% predictive interval.", "Supports burn, hiring, churn, and fundraising timing decisions under uncertainty."],
        ],
        "Table 2. Streamlit application pages and their decision-support purpose.",
        [1.45, 2.55, 2.5],
    )
    add_inline_label_paragraph(
        doc,
        "MS Analytics competencies. ",
        "The project applied data engineering, data validation, exploratory analysis, statistical modeling, forecasting, simulation, model evaluation, visualization, decision-support design, and communication of uncertainty. It also required translating model outputs into software screens that a nontechnical founder could use during a monthly operating review.",
    )

    add_numbered_heading(doc, 4, "Results")
    add_body_paragraph(
        doc,
        f"The validated data show a portfolio with an average SHARF score of {metrics['sharf_score'].mean():.2f} and a median runway of {metrics['runway_months'].median():.2f} months. The mean runway is {metrics['runway_months'].mean():.2f} months, but the median is much lower at {metrics['runway_months'].median():.2f} months, which indicates a right-skewed distribution: a few startups have very long runway while many are closer to cash pressure. The standard deviation of runway is {metrics['runway_months'].std():.2f} months, confirming substantial variation across startup-months.",
    )
    add_report_table(
        doc,
        ["Metric", "Mean", "Median", "Std. dev.", "Interpretation"],
        [
            ["Revenue", f"${metrics['monthly_revenue_usd'].mean():,.0f}", f"${metrics['monthly_revenue_usd'].median():,.0f}", f"${metrics['monthly_revenue_usd'].std():,.0f}", "Traction varies meaningfully across companies and months."],
            ["Cash balance", f"${metrics['cash_balance_usd'].mean():,.0f}", f"${metrics['cash_balance_usd'].median():,.0f}", f"${metrics['cash_balance_usd'].std():,.0f}", "Large spread shows uneven funding positions."],
            ["Monthly burn", f"${metrics['monthly_burn_usd'].mean():,.0f}", f"${metrics['monthly_burn_usd'].median():,.0f}", f"${metrics['monthly_burn_usd'].std():,.0f}", "Burn is operationally material and varies by scale."],
            ["Runway", f"{metrics['runway_months'].mean():.2f}", f"{metrics['runway_months'].median():.2f}", f"{metrics['runway_months'].std():.2f}", "Mean above median signals skew from a few long-runway firms."],
            ["Churn rate", f"{metrics['monthly_churn_rate'].mean():.1%}", f"{metrics['monthly_churn_rate'].median():.1%}", f"{metrics['monthly_churn_rate'].std():.1%}", "Customer-retention risk is visible but not uniform."],
            ["SHARF score", f"{metrics['sharf_score'].mean():.2f}", f"{metrics['sharf_score'].median():.2f}", f"{metrics['sharf_score'].std():.2f}", "Average score sits in the Watch range."],
        ],
        "Table 3. Descriptive statistics for key SHARF operating indicators.",
        [1.25, 1.0, 1.0, 1.0, 2.25],
    )
    add_body_paragraph(
        doc,
        f"Risk-tier counts also show why the application needs decision guidance rather than a passive dashboard. Critical months account for {risk_counts.get('Critical', 0)} of {risk_total} observations ({risk_counts.get('Critical', 0) / risk_total:.1%}), Watch months account for {risk_counts.get('Watch', 0)} observations ({risk_counts.get('Watch', 0) / risk_total:.1%}), and Stable months account for only {risk_counts.get('Stable', 0)} observations ({risk_counts.get('Stable', 0) / risk_total:.1%}). This distribution means most founder-months require monitoring or intervention rather than simple reassurance.",
    )
    add_figure(doc, trend_chart, "Figure 1. Monthly average SHARF score and runway across the validated startup dataset.")
    add_figure(doc, risk_chart, "Figure 2. Distribution of SHARF risk tiers in the validated startup-month dataset.")
    add_body_paragraph(
        doc,
        f"The Monte Carlo forecast layer estimated an average probability of runway falling below six months of {forecast['probability_runway_under_6mo'].mean():.1%}, an average probability of critical risk of {forecast['probability_critical_risk_6mo'].mean():.1%}, and an average probability of cashout of {forecast['probability_cashout_6mo'].mean():.1%}. These probabilities ranged from near zero to 100% across startups, which means the simulated portfolio contains both low-risk and urgent-risk cases. Presenting probability rather than a single deterministic forecast helps the founder understand not just the expected condition, but also the chance of an adverse outcome.",
    )
    add_body_paragraph(
        doc,
        f"The supervised model predicted runway three months ahead. On the temporal holdout, the model achieved MAE={model['metrics']['mae']:.2f} months, RMSE={model['metrics']['rmse']:.2f} months, R^2={model['metrics']['r2']:.3f}, and 90% predictive-interval coverage of {model['metrics']['interval_coverage_90']:.1%}. MAE means the typical absolute runway miss was about four months, while RMSE is slightly higher because larger errors are penalized more strongly. The R^2 value indicates the model explained most of the holdout variation, but the average 90% interval width of {interval_width:.2f} months shows that individual what-if predictions should still be interpreted with caution.",
    )
    add_report_table(
        doc,
        ["Model output", "Value", "Decision meaning"],
        [
            ["MAE", f"{model['metrics']['mae']:.2f} months", "Average absolute error on future holdout months."],
            ["RMSE", f"{model['metrics']['rmse']:.2f} months", "Penalizes larger runway misses more heavily."],
            ["R^2", f"{model['metrics']['r2']:.3f}", "Share of holdout runway variation explained by the model."],
            ["90% interval coverage", f"{model['metrics']['interval_coverage_90']:.1%}", "How often actual runway fell inside the uncertainty interval."],
        ],
        "Table 4. Three-month runway model performance on the temporal holdout.",
        [1.45, 1.35, 3.7],
    )
    add_figure(doc, prediction_chart, "Figure 3. Actual versus predicted runway on the temporal holdout for the three-month runway model.")

    add_numbered_heading(doc, 5, "Discussion and Recommendations")
    add_body_paragraph(
        doc,
        "The main finding is that SHARF can connect data validation, scoring, forecasting, and what-if analysis into a coherent founder workflow. The score gives a quick current-state view, the Monte Carlo forecast explains probabilistic risk, and the supervised model lets the founder test specific operating scenarios. This combination is more useful than a static dashboard because the founder can ask, for example, whether a burn reduction, churn improvement, or fundraising delay meaningfully changes the expected runway and its uncertainty band.",
    )
    add_body_paragraph(
        doc,
        "The user interface is intentionally organized as a decision sequence. Page 1 captures the founder's current assumptions, Page 2 converts those values into a score and risk explanation, Page 3 translates the score into possible operating choices, and Page 4 records the chosen action so the founder can evaluate whether the decision worked. Pages 5 and 6 add predictive accountability: Page 5 shows how the model performed on unseen months before the founder uses it, and Page 6 lets the founder run what-if scenarios with uncertainty displayed next to the point prediction. This sequencing matters because it prevents the application from presenting model output without context.",
    )
    add_body_paragraph(
        doc,
        "The recommended decision rule is to use the lower bound of the 90% predictive interval as the conservative guardrail. If the lower bound is below six months, the founder should prepare a cash-conservation plan or begin fundraising earlier. If the point estimate is acceptable but the interval crosses a risk threshold, the result should be treated as ambiguous and reviewed with additional context. If both the point estimate and lower bound are healthy, the founder can maintain the plan while continuing monthly monitoring.",
    )
    add_body_paragraph(
        doc,
        "Future improvements should replace synthetic operating data with real historical startup records, add model calibration monitoring, expand qualitative founder inputs, and test usability with actual founders or mentors. The predictive model should also be retrained regularly because financing conditions, customer acquisition costs, and hiring markets can shift over time.",
    )

    add_numbered_heading(doc, 6, "Reflection")
    add_body_paragraph(
        doc,
        "This project showed that the hardest part of analytics is not simply building a model; it is deciding what decision the model is supposed to improve. Early versions of SHARF focused on scoring, but the predictive-model requirement made the decision need sharper: founders need to understand whether today's operating choices create enough runway three months from now. That shift improved the application because it connected model performance, uncertainty, and what-if analysis to a real managerial action.",
    )
    add_body_paragraph(
        doc,
        "The project strengthened my ability to move from problem framing to data design, validation, modeling, visualization, and software implementation. It also clarified the importance of humility in analytics. A model can be useful without being absolute, especially when the interface makes uncertainty visible. In relation to the MS Analytics learning objectives, I grew in my ability to combine technical modeling with responsible communication, reproducible code, and decision-centered design.",
    )

    add_numbered_heading(doc, 7, "References")
    references = [
        "Abdullah, D., Erliana, C. I., Bintoro, A., Hartono, Ikhwani, M., & Nazaruddin. (2024). Recipient feasibility decision support system micro small medium business assistance use method Analytic Hierarchy Process and Simple Additives Weighting. International Journal on Informatics Visualization, 8(4), 2119-2124.",
        "Breiman, L. (2001). Statistical modeling: The two cultures. Statistical Science, 16(3), 199-231. https://doi.org/10.1214/ss/1009213726",
        "CB Insights. (2026, March 5). Why startups fail: Top 9 reasons. https://www.cbinsights.com/research/report/startup-failure-reasons-top/",
        "Goodwin, P., Onkal, D., & Thomson, M. E. (2010). Do forecasts expressed as prediction intervals improve production planning decisions? European Journal of Operational Research, 205(1), 195-201. https://doi.org/10.1016/j.ejor.2009.12.020",
        "Hoerl, A. E., & Kennard, R. W. (1970). Ridge regression: Biased estimation for nonorthogonal problems. Technometrics, 12(1), 55-67. https://doi.org/10.1080/00401706.1970.10488634",
        "Mello, L. P., Moraes, G. H. S. M., Fischer, B. B., & Vicentin, D. C. (2026). Pathways to success in equity crowdfunding for sustainable startups: A configurational perspective. Technological Forecasting & Social Change, 222, 124394. https://doi.org/10.1016/j.techfore.2025.124394",
        "Menshikov, V., Ruza, O., Simakhova, A., & Bedianashvili, G. (2025). Common and unique features in the development of startup ecosystems in Latvia, Ukraine, and Georgia. Comparative Economic Research: Central and Eastern Europe, 28(4). https://doi.org/10.18778/1508-2008.28.27",
        "Mseddi, S. (2026). Financing startups and impact investing: Evidence across MENA countries. International Journal of Financial Studies, 14(1), 7. https://doi.org/10.3390/ijfs14010007",
        "Oliva, F. L., Teberga, P. M. F., Testi, L. I. O., Kotabe, M., Del Giudice, M., Kelle, P., & Cunha, M. P. (2022). Risks and critical success factors in the internationalization of born global startups of Industry 4.0: A social, environmental, economic, and institutional analysis. Technological Forecasting & Social Change, 175, 121346. https://doi.org/10.1016/j.techfore.2021.121346",
        "Rady, J., Townsend, D., Hunt, R., & Simpson, J. (2025). The expectations game: The contingent value of hype as a rhetorical strategy in resource mobilization processes among AI startups. Journal of Business Venturing, 40(3), 106499. https://doi.org/10.1016/j.jbusvent.2025.106499",
        "Safari, A., Balicevac Al Ismail, V., Parast, M., Golgeci, I., & Pokharel, S. (2024). Supply chain risk and resilience in startups, SMEs, and large enterprises: A systematic review and directions for research. The International Journal of Logistics Management, 35(2), 680-709. https://doi.org/10.1108/IJLM-10-2022-0422",
        "Ries, E. (2011). The lean startup: How today's entrepreneurs use continuous innovation to create radically successful businesses. Crown Business.",
        "Shmueli, G. (2010). To explain or to predict? Statistical Science, 25(3), 289-310. https://doi.org/10.1214/10-STS330",
        "U.S. Bureau of Labor Statistics. (n.d.). Entrepreneurship and the U.S. economy. https://www.bls.gov/bdm/entrepreneurship/entrepreneurship.htm",
    ]
    for reference in references:
        add_reference(doc, reference)

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    build_document()
