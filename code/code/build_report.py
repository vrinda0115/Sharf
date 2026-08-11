from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = ROOT / "report.docx"
FINAL_DIR = ROOT / "data" / "final"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def style_table(table, widths_dxa: list[int]) -> None:
    table.style = "Table Grid"
    set_table_widths(table, widths_dxa)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.5)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = str(value)
    style_table(table, widths_dxa)


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    doc.add_paragraph(text, style=style)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.1

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(46, 116, 181)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor(46, 116, 181)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(31, 77, 120)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)


def build_report() -> Path:
    profiles = pd.read_csv(FINAL_DIR / "startup_profiles.csv")
    metrics = pd.read_csv(FINAL_DIR / "sharf_monthly_metrics.csv")
    forecast = pd.read_csv(FINAL_DIR / "monte_carlo_risk_forecast.csv")
    benchmark_summary = pd.read_csv(FINAL_DIR / "company_benchmark_summary.csv").iloc[0]
    dictionary = pd.read_csv(FINAL_DIR / "data_dictionary.csv")
    validation = pd.read_csv(FINAL_DIR / "validation_summary.csv")

    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SHARF Data Implementation and Validation Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Startup Health and Risk Forecasting decision-support prototype")
    add_para(doc, "Prepared as the data implementation package for the SHARF monthly founder decision-support tool.")

    doc.add_heading("1. Project continuity and scope check", level=1)
    add_para(doc, "Project title: SHARF Startup Health and Risk Forecasting Decision-Support Tool.")
    add_para(doc, "Decision-maker: A non-technical early-stage startup founder, usually seed to Series A, who needs a monthly operating review without a dedicated finance or data team.")
    add_para(doc, "Problem addressed: Founders often review cash, growth, retention, acquisition efficiency, customer sentiment, hiring capacity, and governance informally and separately. SHARF converts those inputs into a composite health score, risk tier, and decision recommendations.")
    add_para(doc, "Key decisions supported: whether to reduce or maintain burn, whether to approach investors now or prepare/delay, and which operating dimension should receive the next 30-day remediation focus.")
    add_para(doc, "Main metrics and outputs: runway months, month-over-month revenue growth, monthly churn, CAC payback, NPS, headcount, time to fill, governance flag, dimension scores, composite SHARF score, risk tier, burn recommendation, fundraising readiness, remediation priority, and Monte Carlo forecast probabilities for cashout, critical risk, and runway below six months.")
    add_para(doc, "Change from prior assignment: The project scope remains the same, but the risk-forecasting method is refined. Because future startup risk is uncertain and real forecasting data are not known, the implementation now uses Monte Carlo simulation to represent a range of plausible 6-month futures instead of presenting one exact risk forecast. The package also incorporates a sourced company employment/payroll benchmark file provided with the assignment materials. This affects the dataset by adding benchmark and forecast tables, and it affects validation by checking benchmark quality, probability ranges, quantile ordering, simulation counts, and forecast coverage.")

    doc.add_heading("2. Data requirements recap", level=1)
    rows = [
        ["Runway and burn decision", "cash_balance_usd, monthly_burn_usd, runway_months, burn_recommendation", "Generated and derived", "Shows whether spending must be reduced to preserve operating time."],
        ["Revenue trajectory", "monthly_revenue_usd, mom_revenue_growth, revenue_score", "Generated and derived", "Indicates whether growth supports continued burn and investor timing."],
        ["Retention risk", "monthly_churn_rate, retention_score", "Generated and derived", "Captures demand-side weakness that can undermine growth quality."],
        ["Growth efficiency", "customer_acquisition_cost_usd, sales_marketing_spend_usd, new_arr_usd, gross_margin, cac_payback_months", "Generated and derived", "Tests whether acquisition spend is economically sustainable."],
        ["Customer sentiment", "nps, customer_sentiment_score", "Generated", "Adds a customer experience signal before revenue decline appears."],
        ["Team capacity", "headcount, avg_time_to_fill_days, team_capacity_score", "Generated", "Captures whether staffing constraints threaten execution."],
        ["Benchmark calibration", "employment_growth, payroll_growth, payroll_per_employee_usd", "Sourced and derived", "Provides empirical context for employment/payroll volatility used in the Monte Carlo uncertainty layer."],
        ["Governance", "board_or_advisory_group, governance_score", "Manually constructed/generated", "Supports fundraising credibility and accountability review."],
        ["Composite decision logic", "dimension scores, sharf_score, risk_tier, fundraising_readiness, remediation_priority", "Derived", "Converts metrics into the decisions and outputs proposed in the prior blueprint."],
        ["Risk forecast uncertainty", "probability_cashout_6mo, probability_critical_risk_6mo, probability_runway_under_6mo, p10/p50/p90 runway", "Monte Carlo simulated", "Shows risk as probability under uncertainty instead of pretending the future is known exactly."],
    ]
    add_table(doc, ["Component", "Fields needed", "Source plan", "Why it matters"], rows, [1700, 3200, 1700, 2760])

    doc.add_heading("3. Data acquisition and generation process", level=1)
    add_para(doc, "The submitted dataset is synthetic and deterministic. The code uses a fixed random seed in code/generate_sharf_dataset.py, so another person can reproduce the same records by running code/run_pipeline.py.")
    add_bullets(doc, [
        f"The user-provided BQ company employment/payroll CSV.GZ was copied into data/raw/company_employment_payroll_raw.csv.gz and processed into a benchmark table with {int(benchmark_summary['row_count'])} company-year rows across {int(benchmark_summary['company_count'])} companies from {int(benchmark_summary['year_min'])} to {int(benchmark_summary['year_max'])}.",
        "Startup profiles were manually constructed and synthetically generated for 18 ventures across stage, industry, region, founding year, and governance status.",
        "Monthly snapshots were generated for January 2025 through December 2025. Each startup has one row per month.",
        "Revenue, cash, burn, churn, CAC, NPS, headcount, hiring time, gross margin, and macro stress were generated with bounded random variation.",
        "Derived fields were then calculated from the raw monthly observations: runway, month-over-month growth, normalized dimension scores, composite SHARF score, risk tier, and recommendations.",
        "Monte Carlo simulation was added as the risk-forecasting layer. For each startup, the most recent month anchors 1,000 possible 6-month future paths. Revenue growth, burn change, churn, CAC payback, NPS, and hiring time vary using recent historical volatility, bounded assumptions, and employment/payroll volatility from the benchmark dataset.",
        "The original plan's live data integrations were substituted with synthetic data because private startup operating data would not be available, auditable, or ethically appropriate for this assignment."
    ])

    doc.add_heading("4. Empirical basis and assumptions", level=1)
    add_para(doc, "The dataset is not presented as a statistically representative sample of U.S. startups. It is fit-for-purpose prototype data designed to exercise the planned decision logic. Public empirical sources informed the field choices and assumptions: BLS Business Employment Dynamics supports the relevance of establishment survival and business age; FRED demonstrates the availability of macroeconomic context such as interest-rate series; startup fundraising research using Crunchbase supports treating fundraising readiness as stage-sensitive and multi-factor rather than one-metric-only; startup prediction research supports multi-feature backtesting and simulation; and software-startup failure research supports including runway, customer reaction, product-market learning, and execution-capacity signals.")
    add_para(doc, f"The user-provided company benchmark file adds a sourced empirical layer. It contains annual employment and payroll observations, from which the pipeline derives employment growth, payroll growth, and payroll per employee. In the processed benchmark summary, employment growth p10/p50/p90 is {benchmark_summary['employment_growth_p10']}/{benchmark_summary['employment_growth_p50']}/{benchmark_summary['employment_growth_p90']}, payroll growth p10/p50/p90 is {benchmark_summary['payroll_growth_p10']}/{benchmark_summary['payroll_growth_p50']}/{benchmark_summary['payroll_growth_p90']}, and median payroll per employee is ${benchmark_summary['payroll_per_employee_p50']:,.2f}.")
    add_bullets(doc, [
        "Assumption: One startup-month is the correct unit of analysis because the prior SHARF tool is used in monthly founder reviews.",
        "Assumption: Runway is cash divided by monthly burn. This is a transparent operating finance measure and the most direct burn-rate decision input.",
        "Assumption: Higher churn, longer CAC payback, lower NPS, and longer hiring time reduce the health score because they represent weaker retention, acquisition efficiency, customer sentiment, and execution capacity.",
        "Assumption: The AHP-SAW weights are prototype weights, not learned model coefficients. They make the decision logic explainable and auditable.",
        "Assumption: Risk tiers are coarse categories for decision support. They should trigger review and explanation, not deterministic survival claims.",
        "Assumption: Monte Carlo simulation is appropriate for risk forecasting because future revenue, burn, churn, CAC payback, NPS, and hiring conditions are uncertain. The simulation does not discover truth; it makes uncertainty explicit and converts assumptions into auditable probability outputs.",
        "Assumption: Customer and execution fields are included because software-startup failure research links scarce resources, weak learning, negative customer reaction, and flawed business models to startup distress. In this prototype, churn and NPS proxy customer reaction, while time to fill and headcount proxy execution capacity.",
        "Assumption: The annual company benchmark is not startup-specific and does not contain monthly founder operating metrics. It is used only to calibrate employment/payroll uncertainty, not to replace SHARF startup observations."
    ])

    doc.add_heading("5. Final dataset description", level=1)
    risk_counts = metrics["risk_tier"].value_counts().to_dict()
    high_risk_count = int((forecast["monte_carlo_risk_label"] == "High forecast risk").sum())
    moderate_risk_count = int((forecast["monte_carlo_risk_label"] == "Moderate forecast risk").sum())
    lower_risk_count = int((forecast["monte_carlo_risk_label"] == "Lower forecast risk").sum())
    add_para(doc, f"Final structure: five related final data outputs plus a data dictionary and validation outputs. The primary historical SHARF unit of analysis is one startup-month. The final monthly metrics table has {metrics.shape[0]} rows and {metrics.shape[1]} columns. The startup profile table has {profiles.shape[0]} rows and {profiles.shape[1]} columns. The Monte Carlo forecast table has {forecast.shape[0]} rows and {forecast.shape[1]} columns, with one 6-month forecast summary per startup. The sourced benchmark table has {int(benchmark_summary['row_count'])} company-year rows. The SHARF historical time coverage is January 2025 through December 2025, while the benchmark covers {int(benchmark_summary['year_min'])}-{int(benchmark_summary['year_max'])}.")
    add_para(doc, f"Risk-tier coverage in the final dataset: Critical={risk_counts.get('Critical', 0)}, Watch={risk_counts.get('Watch', 0)}, Stable={risk_counts.get('Stable', 0)}. This gives the prototype examples across all planned decision states.")
    add_para(doc, f"Monte Carlo forecast labels: High forecast risk={high_risk_count}, Moderate forecast risk={moderate_risk_count}, Lower forecast risk={lower_risk_count}. The probabilities are based on 1,000 simulated paths per startup over a 6-month horizon.")
    add_para(doc, "Key identifiers: startup_id connects startup_profiles.csv to sharf_monthly_metrics.csv and monte_carlo_risk_forecast.csv. The monthly table grain is enforced by the composite key startup_id plus month. The forecast table grain is one row per startup forecast horizon.")
    dict_rows = dictionary.values.tolist()
    add_table(doc, ["Field", "Definition", "Type"], dict_rows, [2200, 5300, 1860])
    forecast_rows = forecast[
        [
            "startup_id",
            "expected_sharf_score_6mo",
            "probability_cashout_6mo",
            "probability_critical_risk_6mo",
            "probability_runway_under_6mo",
            "monte_carlo_risk_label",
        ]
    ].head(8).values.tolist()
    add_para(doc, "Sample Monte Carlo forecast outputs from the final forecast table:")
    add_table(
        doc,
        ["Startup", "Expected score", "Cashout prob.", "Critical prob.", "Runway <6 prob.", "Forecast label"],
        forecast_rows,
        [1100, 1300, 1350, 1350, 1500, 2760],
    )

    doc.add_heading("6. Data cleaning and transformation", level=1)
    add_para(doc, "The pipeline starts with raw synthetic profile and monthly snapshot files, then produces a scored intermediate table and final submission tables.")
    add_bullets(doc, [
        "Cleaning: Generated values are bounded at creation to prevent impossible negative revenue, negative burn, churn outside 0-1, impossible NPS, and invalid headcount.",
        "Transformations: Runway, month-over-month growth, normalized scores, composite score, and decision outputs are calculated after raw generation.",
        "Benchmark transformation: The sourced annual company file is standardized, numeric employment/payroll fields are parsed, payroll per employee is derived, and annual employment/payroll growth rates are calculated by company.",
        "Monte Carlo transformation: The most recent 6 months of each startup's history estimate uncertainty ranges. The sourced benchmark contributes employment/payroll volatility context. The simulation then produces 1,000 possible 6-month endpoints and summarizes them into probabilities and runway quantiles.",
        "Joins: Startup profile attributes are merged into the monthly table using startup_id.",
        "Missing values: The first month of revenue growth is set to 0 because no prior month exists. Validation confirms no missing cells remain.",
        "Duplicates: The validation script checks uniqueness of startup_id in profiles and startup_id plus month in monthly metrics.",
        "Outliers: Plausible upper and lower bounds are applied to generated fields such as CAC payback, hiring time, churn, NPS, and macro stress."
    ])

    doc.add_heading("7. Data validation", level=1)
    val_rows = validation.assign(passed=validation["passed"].astype(str)).values.tolist()
    add_table(doc, ["Check", "Passed", "Outcome"], val_rows, [2600, 1100, 5660])
    passed_count = int(validation["passed"].astype(str).str.lower().eq("true").sum())
    add_para(doc, f"All {passed_count} validation checks passed. The checks confirm that schemas are complete, identifiers are unique, monthly grain is enforced, all monthly records link to a profile, forecast rows link to profiles, benchmark years and summary counts are valid, no missing cells remain in required final tables, numeric fields are within allowed ranges, the runway formula recomputes correctly, decision outputs are present, planned score inputs are numeric, all startups have 12 months of data, and Monte Carlo forecast outputs have valid probabilities, ordered quantiles, and the expected simulation counts.")

    doc.add_heading("8. Dataset validity justification", level=1)
    add_para(doc, "This dataset should not be treated as universally valid startup-risk data. It is fit for the narrower purpose of building and testing the SHARF course prototype because it matches the specific decision-maker, workflow, metrics, and outputs defined in the earlier blueprint. The intended decision-maker is a non-technical early-stage founder conducting a monthly operating review, and the dataset is structured at exactly that grain: one startup-month.")
    add_para(doc, "For that purpose, the dataset is adequate because it contains the fields needed to compute current health scores, burn recommendations, fundraising readiness, remediation priorities, and 6-month Monte Carlo risk summaries. It also includes enough variation to test Critical, Watch, Stable, and forecast-risk outputs, and it uses the sourced benchmark only where it is relevant: as empirical context for employment/payroll uncertainty rather than as direct startup operating data.")
    add_para(doc, "The dataset is credible for prototype testing because every field is generated, manually constructed, sourced, or derived in an explicit and reproducible way. The code records the generation logic, the final data dictionary documents the fields, and the validation files report the outcome of checks rather than merely proposing checks.")
    add_para(doc, "Important limitations: the startup operating records are synthetic; the sourced benchmark is annual, company-level, and not limited to early-stage startups; the weights, thresholds, and Monte Carlo distributions are assumptions; the macro stress field is a proxy rather than a live economic series; and no real founder behavior or investor outcomes are observed. Because of these limitations, the dataset is useful for testing SHARF's workflow, calculations, uncertainty communication, explanation, and prototype decision logic. It should not be used to make real investment decisions, claim actual survival probabilities, or prove production predictive accuracy.")

    doc.add_heading("9. Reproducibility and submission inventory", level=1)
    inventory_rows = [
        ["report.docx", "Completed assignment report"],
        ["README.txt", "Reproduction instructions and file inventory"],
        ["code/run_pipeline.py", "Runs data generation and validation"],
        ["code/process_company_benchmarks.py", "Processes the user-provided sourced company employment/payroll benchmark"],
        ["code/generate_sharf_dataset.py", "Creates raw, intermediate, and final synthetic SHARF datasets"],
        ["code/monte_carlo_risk_forecast.py", "Runs 1,000 6-month Monte Carlo simulations per startup and summarizes risk probabilities"],
        ["code/validate_sharf_dataset.py", "Validates the final datasets"],
        ["code/build_report.py", "Builds this Word report from final data and validation artifacts"],
        ["data/raw/startup_profiles_raw.csv", "Raw synthetic profile table"],
        ["data/raw/monthly_snapshots_raw.csv", "Raw synthetic monthly observations"],
        ["data/raw/company_employment_payroll_raw.csv.gz", "User-provided sourced company employment/payroll benchmark"],
        ["data/intermediate/monthly_snapshots_scored.csv", "Scored intermediate monthly table"],
        ["data/intermediate/monte_carlo_simulation_paths.csv", "Intermediate endpoint records for all Monte Carlo simulation paths"],
        ["data/final/startup_profiles.csv", "Final startup profile table"],
        ["data/final/sharf_monthly_metrics.csv", "Final startup-month decision-support dataset"],
        ["data/final/monte_carlo_risk_forecast.csv", "Final 6-month probabilistic risk forecast table"],
        ["data/final/company_employment_payroll_benchmark.csv", "Processed sourced company-year employment/payroll benchmark"],
        ["data/final/company_benchmark_summary.csv", "Benchmark summary statistics used by Monte Carlo calibration"],
        ["data/final/data_dictionary.csv", "Field definitions"],
        ["data/final/validation_summary.csv", "Validation check outcomes"],
        ["data/final/validation_results.json", "Machine-readable validation results"],
        ["docs/empirical_basis_and_assumptions.md", "Supporting evidence and assumptions"],
    ]
    add_table(doc, ["File", "Purpose"], inventory_rows, [3300, 6060])
    add_para(doc, "Run order: python code/run_pipeline.py. The report can be rebuilt with python code/build_report.py after the pipeline completes. Required packages are Python, pandas, numpy, and python-docx.")

    doc.add_heading("References", level=1)
    refs = [
        "U.S. Bureau of Labor Statistics. Business Employment Dynamics: Entrepreneurship and the U.S. Economy. https://www.bls.gov/bdm/entrepreneurship/entrepreneurship.htm",
        "Federal Reserve Bank of St. Louis. Federal Funds Effective Rate (FEDFUNDS), FRED. https://fred.stlouisfed.org/series/FEDFUNDS",
        "Gastaud, C., Carniel, T., and Dalle, J.-M. (2019). The varying importance of extrinsic factors in the success of startup fundraising. arXiv. https://arxiv.org/abs/1906.03210",
        "Potanin, M., Chertok, A., Zorin, K., and Shtabtsovsky, C. (2023). Startup success prediction and VC portfolio simulation using CrunchBase data. arXiv. https://arxiv.org/abs/2309.15552",
        "Giardino, C., Wang, X., and Abrahamsson, P. (2017). Why early-stage software startups fail: A behavioral framework. arXiv. https://arxiv.org/abs/1709.04749",
        "Bajwa, S. S., Wang, X., Duc, A. N., and Abrahamsson, P. (2017). Failures to be celebrated: An analysis of major pivots of software startups. arXiv. https://arxiv.org/abs/1710.04037",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    print(build_report())
